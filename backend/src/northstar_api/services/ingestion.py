from __future__ import annotations

import asyncio
import hashlib
import html
import io
import re
import tempfile
import zipfile
from pathlib import Path
from urllib.parse import urlparse
from uuid import UUID

import fitz  # type: ignore[import-untyped]
import structlog
from docx import Document
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from northstar_api.config import get_settings
from northstar_api.models import (
    DocumentChunk,
    IngestionJob,
    JobStatus,
    KnowledgeKind,
    KnowledgeSource,
    KnowledgeStatus,
)
from northstar_api.services.crawler import crawl_website, take_screenshot
from northstar_api.services.llm import nvidia_adapter
from northstar_api.services.object_store import object_store
from northstar_api.services.outbox import enqueue_event
from northstar_api.services.safe_fetch import fetch_public_text
from northstar_api.services.vision import extract_vision_text

logger = structlog.get_logger(__name__)


def semantic_chunks(text: str, target_words: int = 220, overlap_words: int = 35) -> list[str]:
    normalized = re.sub(r"\r\n?", "\n", text).strip()
    if not normalized:
        return []
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", normalized) if part.strip()]
    chunks: list[str] = []
    buffer: list[str] = []
    word_count = 0
    for paragraph in paragraphs:
        paragraph_words = paragraph.split()
        if buffer and word_count + len(paragraph_words) > target_words:
            combined = "\n\n".join(buffer)
            chunks.append(combined)
            overlap = combined.split()[-overlap_words:]
            buffer = [" ".join(overlap)] if overlap else []
            word_count = len(overlap)
        if len(paragraph_words) > target_words * 2:
            for start in range(0, len(paragraph_words), target_words - overlap_words):
                part = paragraph_words[start : start + target_words]
                if part:
                    chunks.append(" ".join(part))
            continue
        buffer.append(paragraph)
        word_count += len(paragraph_words)
    if buffer:
        chunks.append("\n\n".join(buffer))
    return chunks


async def ingest_source_async(session: AsyncSession, source_id: UUID) -> None:
    source = await session.scalar(select(KnowledgeSource).where(KnowledgeSource.id == source_id))
    if not source:
        return
    job = await session.scalar(
        select(IngestionJob)
        .where(IngestionJob.source_id == source.id)
        .order_by(IngestionJob.created_at.desc())
    )
    if job:
        if job.status == JobStatus.SUCCEEDED:
            return
        job.status = JobStatus.RUNNING
        job.step = "chunking"
        job.progress = 20
        job.attempts += 1
    source.status = KnowledgeStatus.PROCESSING
    source.error = None
    await session.commit()
    try:
        content = await _materialize_text(source)
        if not content:
            raise ValueError("Source has no extracted text")
        settings = get_settings()
        if len(content) > settings.knowledge_max_extracted_characters:
            raise ValueError("Source text exceeds the configured extraction limit")
        chunks = semantic_chunks(content)
        if not chunks:
            raise ValueError("Source produced no usable chunks")
        if len(chunks) > settings.knowledge_max_chunks:
            raise ValueError("Source produced more chunks than the configured limit")
        if job:
            job.step = "embedding"
            job.progress = 55
            await session.commit()
        vectors: list[list[float]] = []
        batch_size = 32
        for start in range(0, len(chunks), batch_size):
            vectors.extend(await nvidia_adapter.embed_documents(chunks[start : start + batch_size]))
        await session.execute(delete(DocumentChunk).where(DocumentChunk.source_id == source.id))
        for ordinal, (content_chunk, vector) in enumerate(zip(chunks, vectors, strict=True)):
            session.add(
                DocumentChunk(
                    tenant_id=source.tenant_id,
                    agent_id=source.agent_id,
                    source_id=source.id,
                    ordinal=ordinal,
                    content=content_chunk,
                    token_count=max(1, round(len(content_chunk.split()) * 1.35)),
                    approved=True,
                    revision=source.revision,
                    embedding_model=get_settings().nvidia_embedding_model,
                    embedding=vector,
                    metadata_json={"checksum": hashlib.sha256(content_chunk.encode()).hexdigest()},
                )
            )
        source.status = KnowledgeStatus.READY
        source.chunk_count = len(chunks)
        source.size_label = f"{len(content.split()):,} words"
        if job:
            job.status = JobStatus.SUCCEEDED
            job.step = "complete"
            job.progress = 100
        enqueue_event(
            session,
            tenant_id=source.tenant_id,
            aggregate_type="knowledge_source",
            aggregate_id=source.id,
            event_type="knowledge.source.ingested.v1",
            payload={"sourceId": str(source.id), "agentId": str(source.agent_id), "chunks": len(chunks)},
        )
        await session.commit()
    except Exception as exc:
        source.status = KnowledgeStatus.FAILED
        source.error = "Ingestion failed; inspect the job log."
        if job:
            job.status = JobStatus.FAILED
            job.step = "failed"
            job.error_json = {"type": type(exc).__name__, "message": str(exc)[:400]}
        await session.commit()
        raise


async def _materialize_text(source: KnowledgeSource) -> str:
    if source.content:
        return source.content.strip()
    if source.url:
        if source.kind == KnowledgeKind.SITEMAP:
            return await _materialize_sitemap(source)

        if source.kind == KnowledgeKind.URL:
            return await _materialize_webpage_with_vision(source)

        text, final_url = await fetch_public_text(source.url)
        source.url = final_url
        return text.strip()
    if source.object_key:
        data, metadata = await object_store.download(source.object_key)
        content_type = str(metadata.get("content_type") or "").lower()
        if content_type == "application/pdf" or source.name.lower().endswith(".pdf"):
            reader = PdfReader(io.BytesIO(data))
            settings = get_settings()

            if len(reader.pages) > settings.knowledge_max_pdf_pages:
                raise ValueError("PDF exceeds the configured page limit")

            pages: list[str] = []

            # Existing PDF text extraction
            extracted_characters = 0
            for page in reader.pages:
                text = page.extract_text() or ""
                extracted_characters += len(text)

                if extracted_characters > settings.knowledge_max_extracted_characters:
                    raise ValueError("PDF text exceeds the configured extraction limit")

                if text.strip():
                    pages.append(text.strip())

            # NEXORA Vision extraction
            try:
                with tempfile.TemporaryDirectory() as temp_dir:
                    pdf = fitz.open(stream=data, filetype="pdf")

                    for page_number in range(len(pdf)):
                        page = pdf[page_number]

                        # Render PDF page as PNG
                        pixmap = page.get_pixmap(
                            matrix=fitz.Matrix(2, 2),
                            alpha=False,
                        )

                        image_path = Path(temp_dir) / f"vision_page_{page_number + 1:03d}.png"
                        pixmap.save(str(image_path))

                        # Extract visible text using Vision LLM
                        vision_text = extract_vision_text(str(image_path))

                        if vision_text.strip():
                            pages.append(
                                f"VISUAL CONTENT FROM PDF PAGE {page_number + 1}:\n{vision_text.strip()}"
                            )

                    pdf.close()

            except Exception as exc:
                logger.warning(
                    "pdf_vision_extraction_failed",
                    source_id=str(source.id),
                    error=type(exc).__name__,
                )

            return "\n\n".join(pages).strip()
        if "wordprocessingml.document" in content_type or source.name.lower().endswith(".docx"):
            settings = get_settings()

            try:
                with zipfile.ZipFile(io.BytesIO(data)) as archive:
                    members = archive.infolist()
                    expanded_size = sum(member.file_size for member in members)

                    if len(members) > 5_000:
                        raise ValueError("DOCX contains too many archive members")

                    if expanded_size > settings.knowledge_max_docx_uncompressed_bytes:
                        raise ValueError("DOCX exceeds the configured expanded-size limit")

            except zipfile.BadZipFile as exc:
                raise ValueError("DOCX archive is invalid") from exc

            document = Document(io.BytesIO(data))

            sections: list[str] = []

            # Existing paragraph extraction
            paragraphs = [
                paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
            ]

            if paragraphs:
                sections.append("\n\n".join(paragraphs))

            # Existing document tables + structured table text
            tables: list[str] = []

            for table in document.tables:
                rows: list[str] = []

                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]

                    if any(cells):
                        rows.append(" | ".join(cells))

                if rows:
                    tables.append("\n".join(rows))

            if tables:
                sections.append("TABLE CONTENT FROM DOCX:\n\n" + "\n\n".join(tables))

            # NEXORA Vision extraction for embedded images
            try:
                with tempfile.TemporaryDirectory() as temp_dir:
                    with zipfile.ZipFile(io.BytesIO(data)) as archive:
                        image_members = [
                            member
                            for member in archive.infolist()
                            if member.filename.startswith("word/media/")
                        ]

                        for image_number, member in enumerate(
                            image_members,
                            start=1,
                        ):
                            suffix = Path(member.filename).suffix or ".png"

                            image_path = Path(temp_dir) / f"docx_image_{image_number:03d}{suffix}"

                            with archive.open(member) as source_file:
                                image_path.write_bytes(source_file.read())

                            vision_text = await asyncio.to_thread(
                                extract_vision_text,
                                str(image_path),
                            )

                            if vision_text.strip():
                                sections.append(
                                    f"VISUAL CONTENT FROM DOCX IMAGE {image_number}:\n{vision_text.strip()}"
                                )

            except Exception as exc:
                logger.warning(
                    "docx_vision_extraction_failed",
                    source_id=str(source.id),
                    error=type(exc).__name__,
                )

            return "\n\n".join(sections).strip()
        if content_type.startswith("text/") or source.name.lower().endswith((".txt", ".md")):
            return data.decode("utf-8", errors="replace").strip()
        raise ValueError("Uploaded file type is not supported")
    return ""


async def _materialize_webpage_with_vision(source: KnowledgeSource) -> str:
    """
    Crawl a website, extract normal HTML text, take screenshots,
    and add NEXORA Vision text before the existing Northstar
    chunking and embedding pipeline.
    """
    assert source.url is not None

    settings = get_settings()

    initial_text, final_url = await fetch_public_text(source.url)
    source.url = final_url

    try:
        # Run the NEXORA crawler without blocking the async ingestion loop.
        pages = await asyncio.to_thread(
            crawl_website,
            final_url,
            10,
        )
    except Exception as exc:
        logger.warning(
            "website_crawl_failed",
            source_id=str(source.id),
            error=type(exc).__name__,
        )
        return initial_text.strip()

    sections: list[str] = []

    for page_number, page in enumerate(pages, start=1):
        page_url = str(page.get("url") or "")
        page_text = str(page.get("text") or "").strip()

        if not page_url:
            continue

        section_parts: list[str] = []

        if page_text:
            section_parts.append(page_text)

        # NEXORA screenshot + Vision extraction.
        try:
            screenshot_path = await asyncio.to_thread(
                take_screenshot,
                page_url,
                f"web_{source.id}_{page_number:03d}.png",
            )

            vision_text = await asyncio.to_thread(
                extract_vision_text,
                screenshot_path,
            )

            if vision_text.strip():
                section_parts.append(f"VISUAL CONTENT FROM WEB PAGE:\n{vision_text.strip()}")

        except Exception as exc:
            logger.warning(
                "website_vision_extraction_failed",
                source_id=str(source.id),
                page_url=page_url,
                error=type(exc).__name__,
            )

        if section_parts:
            sections.append(f"Source page: {page_url}\n\n" + "\n\n".join(section_parts))

    if not sections:
        return initial_text.strip()

    combined = "\n\n---\n\n".join(sections)

    return combined[: settings.knowledge_max_extracted_characters].strip()


async def _materialize_sitemap(source: KnowledgeSource) -> str:
    assert source.url is not None
    settings = get_settings()
    sitemap, final_sitemap_url = await fetch_public_text(source.url, max_bytes=1_000_000)
    source.url = final_sitemap_url
    sitemap_host = (urlparse(final_sitemap_url).hostname or "").lower()
    raw_locations = re.findall(
        r"<(?:[A-Za-z0-9_-]+:)?loc\b[^>]*>(.*?)</(?:[A-Za-z0-9_-]+:)?loc>",
        sitemap,
        flags=re.IGNORECASE | re.DOTALL,
    )
    locations: list[str] = []
    for raw_location in raw_locations:
        location = html.unescape(raw_location).strip()
        parsed = urlparse(location)
        if (
            parsed.scheme.lower() not in {"http", "https"}
            or (parsed.hostname or "").lower() != sitemap_host
            or location in locations
        ):
            continue
        locations.append(location)
        if len(locations) >= settings.knowledge_sitemap_max_urls:
            break
    if not locations:
        raise ValueError("Sitemap did not contain same-host HTTP(S) page URLs")

    semaphore = asyncio.Semaphore(5)

    async def fetch_page(page_url: str) -> tuple[str, str] | None:
        try:
            async with semaphore:
                page_text, final_url = await fetch_public_text(page_url, max_bytes=500_000)
            return final_url, page_text
        except Exception as exc:
            logger.warning(
                "sitemap_page_fetch_failed",
                host=(urlparse(page_url).hostname or "")[:255],
                error=type(exc).__name__,
            )
            return None

    pages = await asyncio.gather(*(fetch_page(location) for location in locations))
    sections: list[str] = []
    remaining = settings.knowledge_max_extracted_characters
    for page in pages:
        if not page or not page[1].strip() or remaining <= 0:
            continue
        final_url, page_text = page
        section = f"Source page: {final_url}\n\n{page_text.strip()}"
        section = section[:remaining]
        sections.append(section)
        remaining -= len(section)
    if not sections:
        raise ValueError("Sitemap pages did not yield usable text")
    return "\n\n---\n\n".join(sections)
